from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent


def set_run(run, size=11, bold=False, color=None, name='Calibri'):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, size=18, bold=True, color=(0, 51, 102))
    p.space_after = Pt(12)


def add_heading1(doc, text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    set_run(run, size=16, bold=True, color=(0, 51, 102))
    p.space_before = Pt(16)
    p.space_after = Pt(8)


def add_heading2(doc, text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    set_run(run, size=14, bold=True, color=(0, 51, 102))
    p.space_before = Pt(12)
    p.space_after = Pt(6)


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=bold)
    p.space_after = Pt(6)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_run(run)
    p.space_after = Pt(4)


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    set_run(run)
    p.space_after = Pt(4)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(6)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, bold=True, color=(133, 100, 4))
    p2 = cell.add_paragraph(body)
    set_run(p2.runs[0] if p2.runs else p2.add_run(body))
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFF3CD')
    cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        set_run(r, bold=True, color=(255, 255, 255))
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '003366')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            set_run(r)
    doc.add_paragraph()


def build_zavuch_manual():
    doc = Document()
    add_title(doc, 'Дастури истифодабарии Портали ягонаи маорифи ноҳияи Зафаробод')
    add_title(doc, '(Барои ҳамоҳангсозон, муовинони директор ва омӯзгорон)')
    add_paragraph(doc, 'Ин дастур ба муовинони директор (завучҳо), директорони муассиса ва омӯзгорон дастурҳоро оид ба истифодаи Портали ягонаи маорифи ноҳияи Зафаробод пешниҳод менамояд. Дар ин ҷо қадам-ба-қадам шарҳ дода шудааст, ки чӣ тавр ба тизим ворид шудан, синфҳоро идора кардан, хонандагон ва омӯзгоронро тезкор ворид намудан, дарсҳоро тақсим кардан ва журналҳои холгузорӣ пур кардан.')

    add_heading1(doc, '1. Вуруд ба тизим')
    add_paragraph(doc, 'Порталро танҳо корбарони бақайдгирифта метавонанд истифода баранд. Завучҳо ҳар як муассисаро идора мекунанд, директорҳо маълумоти муассисаи худ ва омӯзгоронро мебинанд, омӯзгорон холгузорӣ мекунанд.')
    add_numbered(doc, 'Браузери худро (Chrome, Edge ё Firefox) кушода, суроғаи порталро нависед.')
    add_numbered(doc, 'Дар саҳифаи вуруд логин ва рамзи худро ворид кунед.')
    add_numbered(doc, 'Логинҳои завуч ба шакли zavuch_[рақами мактаб] мебошанд, масалан zavuch_20 барои муассисаи №20.')
    add_numbered(doc, 'Рамз ҳангоми бақайдгирӣ дода шудааст. Агар онро фаромӯш карда бошед, бо Оҳсар Дӯсарович тамос гиред.')
    add_table(doc, ['Рол', 'Логин (намуна)', 'Рамз (намуна)'], [
        ['Маориф мудир', 'director_20', 'Director_20_2026@'],
        ['Завуч / Ҳамоҳангсоз', 'zavuch_20', 'Zavuch_20_2026@'],
        ['Омӯзгор', 'teacher_M20_1', 'Teacher_M20_1@2026'],
    ])
    add_note(doc, 'Эзоҳ:', 'Логин ва рамзро бо дигарон нашарик намоед. Дар сурати шубҳа дар бехатарии ҳисоб, дарҳол онро иваз кунед.')

    add_heading1(doc, '2. Идоракунии синфҳо')
    add_paragraph(doc, 'Пас аз вуруд, аз менюи асосӣ «Муассисаҳо»-ро зер карда, муассисаи худро интихоб намоед. Сипас «Синфҳо»-ро пахш кунед.')
    add_bullet(doc, 'Саҳифаи «Синфҳо» рӯйхати синфҳои муассисаро нишон медиҳад.')
    add_bullet(doc, 'Ҳар як синф аз рӯи хонандагоне, ки ворид шудаанд, ба таври худкор пайдо мешавад.')
    add_bullet(doc, 'Барои дидани рӯйхати хонандагони синф, дар сатри он синф тугмаи «Холҳо»-ро зер кунед.')
    add_bullet(doc, 'Барои иловаи синфи нав, аввал хонандагони он синфро ё дастӣ ворид кунед, ё шаблони Excel-ро истифода баред.')
    add_bullet(doc, 'Синфҳои нав ҳангоми воридкунии аввалин хонанда аз Excel ё дастӣ ба таври худкор пайдо мешаванд.')

    add_heading1(doc, '3. Вориди якбораи хонандагон тавассути Excel')
    add_paragraph(doc, 'Барои суръат бахшидани вориди рӯйхати хонандагон, портал шаблони оқилонаи Excel дорад. Ин шаблон рақами хонандаро ба таври худкор месозад.')
    add_heading2(doc, '3.1. Боргирии шаблон')
    add_numbered(doc, 'Аз саҳифаи «Синфҳо» ё аз саҳифаи ҷузъиёти синф, тугмаи «Боргирии шаблон»-ро зер кунед.')
    add_numbered(doc, 'Файли «Шаблони_Хонандагон.xlsx» ба компютератон бор гирифта мешавад.')
    add_numbered(doc, 'Шаблонро бо Microsoft Excel ё LibreOffice Calc кушоед.')
    add_heading2(doc, '3.2. Пур кардани шаблон')
    add_bullet(doc, 'Сутуни A — «№ мактаб»: дорои формула аст, онро тағйир надиҳед. Сутун қулф карда шудааст.')
    add_bullet(doc, 'Сутуни B — «№ синф»: дорои формула аст, онро тағйир надиҳед. Сутун қулф карда шудааст.')
    add_bullet(doc, 'Сутуни C — «Ному насаб»: номи хонандаро бо ҳарфи калони аввал нависед, масалан «Раҷабов А.». Сутун боз аст.')
    add_bullet(doc, 'Сутуни D — «Синф»: синфи хонандаро нависед, масалан «11-А», «10-Б», «5-В». Сутуни матнӣ аст.')
    add_bullet(doc, 'Формат рақами телефон ё рақамҳои дигар: агар қимматҳои рақамӣ ворид кунед, Excel метавонад онҳоро ба рақам табдил диҳад. Барои пешгирии ин, сутунро ҳамчун матн формат кунед ё апостроф гузоред.')
    add_table(doc, ['Сутун', 'Ном', 'Ҳолат', 'Шарҳ'], [
        ['A', '№ мактаб', 'Қулф кардашуда', 'Формулаи худкор, дастакалӣ тағйир дода намешавад'],
        ['B', '№ синф', 'Қулф кардашуда', 'Формулаи худкор барои рақами синф'],
        ['C', 'Ному насаб', 'Боз', 'Номи пурраи хонанда'],
        ['D', 'Синф', 'Боз', 'Синф бо ҳарф, масалан 11-А'],
    ])
    add_heading2(doc, '3.3. Боркунии файл ба портал')
    add_numbered(doc, 'Файлро нигоҳ дошта, портал кушоед.')
    add_numbered(doc, 'Ба саҳифаи «Синфҳо» баргашта, тугмаи «Воридоти Excel»-ро зер кунед.')
    add_numbered(doc, 'Файли пуршударо интихоб кунед ва бор кунед.')
    add_numbered(doc, 'Пас аз боркунии муваффақ, хонандагон дар рӯйхати синф пайдо мешаванд.')
    add_note(doc, 'Эзоҳ:', 'Агар сутунҳои A ё B иваз карда шаванд, формулаҳо вайрон мешаванд. Парол барои қуфл: maorif_zafarobod.')

    add_heading1(doc, '4. Вориди якбораи омӯзгорон тавассути Excel')
    add_paragraph(doc, 'Барои суръат бахшидани вориди омӯзгорон, аз шаблони махсус истифода баред.')
    add_numbered(doc, 'Аз менюи «Муассисаҳо» муассисаи худро кушоед.')
    add_numbered(doc, 'Ба «Омӯзгорон» гузаред.')
    add_numbered(doc, 'Тугмаи «Шаблони омӯзгорон»-ро зер кунед. Файли «Shabloni_Omuzgoron.xlsx» бор гирифта мешавад.')
    add_numbered(doc, 'Сутуни A дорои формулаи рақами худкор аст. Сутунҳои B, C, D ва E бозанд.')
    add_table(doc, ['Сутун', 'Ном', 'Ҳолат', 'Шарҳ'], [
        ['A', '№', 'Қулф кардашуда', 'Формулаи рақами худкор'],
        ['B', 'Ному насаби омӯзгор', 'Боз', 'Номи пурра'],
        ['C', 'Рақами телефон', 'Боз, матн', 'Масалан +992901234567'],
        ['D', 'Маълумот', 'Боз', 'Дараҷаи таълимӣ'],
        ['E', 'Ихтисос', 'Боз', 'Фан ё ихтисос'],
    ])
    add_numbered(doc, 'Файлро пур карда, нигоҳ дошта, тавассути тугмаи «Воридоти омӯзгорон» бор кунед.')
    add_numbered(doc, 'Пас аз ворид, логин ва рамзи ҳар омӯзгор намоиш дода мешавад. Ин рӯйхатро чоп кунед ё нигоҳ доред.')
    add_note(doc, 'Маслиҳат:', 'Рақами телефонро ҳамчун матн ворид кунед, то Excel онро рақам накунад. Номро бо ҳарфи калони аввал нависед.')

    add_heading1(doc, '5. Тақсимоти дарсҳо')
    add_paragraph(doc, 'Ин қисм имкон медиҳад, ки барои ҳар як синф ва фан, омӯзгори масъул таъин карда шавад.')
    add_numbered(doc, 'Аз менюи боло «Тақсимоти дарсҳо»-ро интихоб кунед.')
    add_numbered(doc, 'Дар ҷадвал дар ҳар сатр синф, фан ва рӯйхати омӯзгорон нишон дода мешавад.')
    add_numbered(doc, 'Аз менюи кушодашаванда дар сутуни «Омӯзгори масъул» омӯзгори дилхоҳро интихоб кунед.')
    add_numbered(doc, 'Барои ҳар як сатри дигар, ки бояд таъин шавад, ҳамин амалро такрор кунед.')
    add_numbered(doc, 'Тугмаи «Сабт кардан»-ро зер кунед. Пас аз он паёми муваффақият намоиш дода мешавад.')
    add_note(doc, 'Эзоҳ:', 'Танҳо омӯзгорони муассисаи худатон дар рӯйхати интихоб нишон дода мешаванд.')

    add_heading1(doc, '6. Холгузорӣ барои омӯзгорон')
    add_paragraph(doc, 'Омӯзгорон пас аз вуруд метавонанд журнали холгузории худро кушоянд.')
    add_numbered(doc, 'Омӯзгор бо логини худ (teacher_M[рақами мактаб]_[counter]) ворид мешавад, масалан teacher_M20_1.')
    add_numbered(doc, 'Синф ва фани вобастаро аз рӯйхати худ интихоб мекунад.')
    add_numbered(doc, 'Дар саҳифаи журнал:')
    add_bullet(doc, 'Барои ҳар як хонанда холи рақамӣ (1-10) ворид кунед.')
    add_bullet(doc, 'Ҳозирӣ (+ ё -) ва хулқ-атвор (1-5) сабт кунед.')
    add_bullet(doc, 'Санаи холгузориро тасдиқ кунед.')
    add_numbered(doc, 'Холҳои чорякро метавонед ба таври худкор ё дастӣ ҳисоб карда, ба низом ворид намоед.')
    add_note(doc, 'Эзоҳ:', 'Холҳо ба таври онлайн нигоҳ дошта мешаванд ва зуд ба рейтинги синф ва муассиса таъсир мерасонанд.')

    path = BASE_DIR / 'Дастури_истифодабарии_Портали_Маориф_барои_Завуч.docx'
    doc.save(path)
    print('Saved Tajik manual')


def build_admin_manual():
    doc = Document()
    add_title(doc, 'Веб-Портал: Йўриқномаи умумий барои Администратор ва Маориф мудири')
    add_title(doc, '(Зафаробод тумани, ягона портали маориф)')
    add_paragraph(doc, 'Ушбу йўриқнома Зафаробод туманидаги маориф бўлимининг ягона веб-порталини бошқариш, созлаш ва хавфсиз ишлатиш учун тайёрланди. Унда техник тузилма, администраторлик ваколатлари, GitHub Desktop ва автоматик деплой воситалари ҳақида батафсил маълумот берилган.')

    add_heading1(doc, '1. Сервер ва лойиҳа тузилмаси')
    add_paragraph(doc, 'Портал Linux серверда, Gunicorn + Nginx технологиялари ёрдамида ишлайди. Маълумотлар базаси SQLite (db.sqlite3) шаклида сақланади.')
    add_bullet(doc, 'Маҳаллий ишлаб чиқиш каталоги: D:/Loihalar/Maorif_Portal')
    add_bullet(doc, 'Ишлаб чиқиш мухити: Python 3.12 + Django 4.2')
    add_bullet(doc, 'Продакшн сервер: 169.58.154.191')
    add_bullet(doc, 'Продакшн каталог: /var/www/maorif_zafarobod/')
    add_bullet(doc, 'Продакшн лойиҳа: /var/www/maorif_zafarobod/maorif_portal/')
    add_bullet(doc, 'Gunicorn: /var/www/maorif_zafarobod/venv/bin/gunicorn')
    add_bullet(doc, 'Nginx конфигурацияси: /etc/nginx/sites-available/maorif')
    add_bullet(doc, 'База: /var/www/maorif_zafarobod/db.sqlite3')
    add_bullet(doc, 'Захира нусхалари: /root/maorif_backups_<sana>')
    add_table(doc, ['Амал', 'Буйруқ'], [
        ['Тизимни текшириш', 'python manage.py check'],
        ['Миграцияларни яратиш', 'python manage.py makemigrations'],
        ['Миграцияларни қўллаш', 'python manage.py migrate'],
        ['Статик файлларни йиғиш', 'python manage.py collectstatic --noinput'],
        ['Серверни ишга тушириш', '.\\venv\\Scripts\\python manage.py runserver'],
    ])
    add_note(doc, 'Эслатма:', 'Продакшн серверда энг охирги ўзгартиришларни қайта ишга тушириш учун Gunicorn ва Nginx-ни қайта ишга тушириш тавсия этилади.')

    add_heading1(doc, '2. Маориф мудири учун умумий метрикалар')
    add_paragraph(doc, 'Маориф бўлими мудири порталда умумий рейтинг ва статистикаларни кўриши мумкин. Ҳозирда портал 31 та муассисани бирлаштиради.')
    add_bullet(doc, 'Асосий панелда умумий миқдорлар: муассисалар, хонандалар, омӯзгорлар ва нисбат кўрсатилган.')
    add_bullet(doc, 'Рейтинги муассисаҳо: ҳар бир мактабнинг холлари миёна.')
    add_bullet(doc, 'Рейтинги фанҳо: ҳар бир фан бўйича туман миқёсида хол.')
    add_bullet(doc, 'Рейтинги синфҳо: синфлар бўйича ноҳиявий ва муассисаи ҷой.')
    add_bullet(doc, 'Таносуби омӯзгор ба хонанда ҳисобланиб, панелда кўрсатилади.')
    add_table(doc, ['Кўрсаткич', 'Маъноси'], [
        ['Муассисаҳо', 'Тумандаги умумий мактаб / литсей / томактаб миқдори'],
        ['Хонандагон', 'Барча муассисалардаги умумий хонанда сони'],
        ['Кормандон', 'Омӯзгор ва кадрлар сони'],
        ['Таносуб', 'Хонанда / омӯзгор нисбати'],
    ])
    add_note(doc, 'Эслатма:', 'Рейтинглар кунлик холгузорилар ва чораклик холларга асосланиб автоматик ҳисобланади.')

    add_heading1(doc, '3. Администраторлик ваколатлари ва Админ Панел')
    add_paragraph(doc, 'Портални техник ривожлантирувчи ва администратор Оҳсар Дӯсарович бошқариб туриши лозим. Админ панел орқали коди ёзмасдан ҳамма маълумотларни бошқариш мумкин.')
    add_numbered(doc, 'Браузерда /admin/ манзилига кириш.')
    add_numbered(doc, 'Superuser ҳисоби ёки администратор логини билан кириш.')
    add_numbered(doc, 'Муассисалар рўйхатини қўшиш, таҳрирлаш ва ўчириш.')
    add_numbered(doc, 'Корбарлар, UserProfile ва TeacherProfile ёзувларини бошқариш.')
    add_numbered(doc, 'Синфлар, фанлар, хонандалар ва омӯзгорларни қўллаб-қувватлаш ёки тузатиш.')
    add_numbered(doc, 'Дарс тақсимоти, чорак баҳолари ва хулқ-атвор маълумотларини назорат қилиш.')
    add_numbered(doc, 'Агар маълумотлар базасида хатолик юз берса, SQLite файлини ёки захира нусхасини тиклаш.')
    add_table(doc, ['Объект', 'Бошқариш имконияти'], [
        ['School', 'Муассиса номи, тури, директор, телефон'],
        ['User / UserProfile', 'Логин, рамз, рол ва муассиса'],
        ['TeacherProfile', 'Омӯзгор профили ва ихтисос'],
        ['ClassSubject', 'Синф-фан ва омӯзгор алоқалари'],
        ['Student / Grade', 'Хонандар, холлар, давомат'],
    ])
    add_note(doc, 'Хавфсизлик:', 'Админ панелга фақат ваколатли шахслар кириши керак. Муҳим амаллардан олдин базани захиралаш тавсия этилади.')

    add_heading1(doc, '4. GitHub Desktop ва deploy_to_server.py')
    add_paragraph(doc, 'Янги шаблонлар, кўрсатмалар ёгидаги ўзгартиришларни продакшн серверга чиқариш учун GitHub Desktop ва маҳаллий deploy_to_server.py скриптидан фойдаланиш кўрсатилган.')
    add_heading2(doc, '4.1. GitHub Desktop орқали қадамлар')
    add_numbered(doc, 'GitHub Desktop-ни очиб, Maorif_Portal лойиҳасини танланг.')
    add_numbered(doc, 'Ғайри зарур файллар (.pyc, venv) учун .gitignore мавжудлигини текширинг.')
    add_numbered(doc, 'Summary майдонида қисқа commit хабарини ёзинг.')
    add_numbered(doc, 'Commit to main тугмаси билан маҳаллий commit яратинг.')
    add_numbered(doc, 'Push origin main тугмаси билан ўзгартиришларни GitHub-га юборинг.')
    add_heading2(doc, '4.2. Автоматик деплой')
    add_numbered(doc, 'deploy_to_server.py скрипти маҳаллий венвда paramiko ёрдамида серверга уланади.')
    add_numbered(doc, 'Скрипт сервердаги /var/www/maorif_zafarobod/ каталогида origin/main-ни қайта чиқариб олади.')
    add_numbered(doc, 'db.sqlite3 ва schools.db файлларини захиралаб, янги кодни чиқариб, қайта тиклайди.')
    add_numbered(doc, 'python manage.py migrate, Gunicorn ва Nginx-ни қайта ишга туширади.')
    add_table(doc, ['Файл', 'Вазифаси'], [
        ['.gitignore', 'venv ва __pycache__ гитда сақланмаслиги учун'],
        ['deploy_to_server.py', 'SSH орқали автоматик продакшн янгилаш'],
        ['generate_manuals.py', 'Кўрсатма .docx ҳужжатларини яратиш'],
    ])
    add_note(doc, 'Эслатма:', 'deploy_to_server.py паролни getpass билан қабул қилади. SSH_PASSWORD ёки REPO_URL муҳит ўзгарувчилари орқали автоматлаштириш мумкин.')

    add_heading1(doc, '5. Хавфсизлик ва Заҳиралаш')
    add_paragraph(doc, 'Лойиҳа файларини вақти-вақти билан таққослаш ва захиралаш зарур.')
    add_bullet(doc, 'GitHub Desktop ёки Git маҳаллий орнатилганлигини текширинг.')
    add_bullet(doc, 'Loyiҳа каталогида git status натижасини кўриб, ўзгартирилган файлларни тасдиқланг.')
    add_bullet(doc, 'Ҳар бир муҳим ўзгаришдан кейин commit қилинг.')
    add_bullet(doc, 'db.sqlite3, schools.db ва media каталогини ҳам вақти-вақти билан нусха олиб туриш тавсия этилади.')
    add_bullet(doc, 'Сирли маълумотлар (пароллар, калитлар) гитга тарқалмаслиги учун .gitignore файлини текширинг.')
    add_table(doc, ['Амал', 'Буйруқ ёки ҳаракат'], [
        ['Ўзгаришларни кўриш', 'git status'],
        ['Индексация', 'git add .'],
        ['Сақлаш', 'git commit -m "\u042f\u043d\u0433\u0438 \u045e\u0437\u0433\u0430\u0440\u0442\u0438\u0448\u043b\u0430\u0440"'],
        ['Юбориш', 'git push origin main'],
        ['Деплой', 'python deploy_to_server.py'],
    ])
    add_note(doc, 'Эслатма:', 'Сирли маълумотлар (пароллар, калитлар) гитга тарқалмаслиги учун .gitignore файлини текширинг.')

    path = BASE_DIR / 'Веб_Портал_Админ_Умумий_Йўриқномаси.docx'
    doc.save(path)
    print('Saved Uzbek Cyrillic admin manual')


if __name__ == '__main__':
    build_zavuch_manual()
    build_admin_manual()
